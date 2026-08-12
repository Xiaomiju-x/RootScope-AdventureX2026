from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "AdventureX_RootScope_固定式根区灌溉舱_最终方案.md"
IMAGE = ROOT / "output" / "rootscope" / "images" / "RootScope_固定式根区灌溉舱_产品模型图_三泵版.png"
OUTPUT = ROOT / "output" / "rootscope" / "AdventureX_RootScope_固定式根区灌溉舱_最终方案.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16324F"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F4F6F9"
MID_GRAY = "D7DEE7"
TEXT = "243342"
MUTED = "5E6B78"
ORANGE = "E8862C"
WHITE = "FFFFFF"
RED = "C0392B"
GREEN = "2D7D46"

PAGE_WIDTH = 12240   # Letter, 8.5 in
PAGE_HEIGHT = 15840  # Letter, 11 in
MARGIN = 1440        # 1 in
BODY_WIDTH = 9360


def set_east_asia(run, latin: str = "Calibri", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], total_width: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_paragraph_keep(paragraph, keep_with_next=True, keep_together=True)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header_and_rows(table) -> None:
    repeat_header(table.rows[0])
    for row in table.rows:
        keep_row_together(row)


def set_paragraph_keep(paragraph, keep_with_next: bool = False, keep_together: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_together:
        p_pr.append(OxmlElement("w:keepLines"))
    p_pr.append(OxmlElement("w:widowControl"))


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fld_run = OxmlElement("w:r")
    fld_rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    fld_rpr.extend([color, size])
    fld_run.append(fld_rpr)
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_run.append(fld_text)
    fld.append(fld_run)
    paragraph._p.append(fld)


def style_section(section, landscape: bool = False, compact: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        margin = Inches(0.65 if compact else 0.75)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        margin = Inches(1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def install_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_specs = {
        "Title": (28, NAVY, 0, 10),
        "Subtitle": (15, BLUE, 0, 12),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(TEXT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.extend([fonts, color, underline])
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


TOKEN_RE = re.compile(r"(\[[^\]]+\]\([^\)]+\)|`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text: str, *, size: float | None = None, color: str | None = None, bold: bool = False) -> None:
    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.bold = bold
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
            set_east_asia(run)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(size or 9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            set_east_asia(run, latin="Consolas")
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
            set_east_asia(run)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        set_east_asia(run)


def add_header_footer(section, first: bool = False) -> None:
    section.different_first_page_header_footer = first
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("ROOTSCOPE  /  ADVENTUREX 2026")
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor.from_string(NAVY)
    set_east_asia(r1)
    r2 = p.add_run("    固定式根区灌溉舱 · 可见、可量、可追溯")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia(r2)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    add_page_field(fp)

    if first:
        first_header = section.first_page_header
        first_header.is_linked_to_previous = False
        first_header.paragraphs[0].clear()
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        first_footer.paragraphs[0].clear()


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    style_section(section)
    add_header_footer(section, first=True)

    accent = doc.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent.autofit = False
    set_table_geometry(accent, [800, 8560], BODY_WIDTH)
    set_cell_shading(accent.cell(0, 0), ORANGE)
    set_cell_shading(accent.cell(0, 1), NAVY)
    for cell in accent.rows[0].cells:
        cell.height = Inches(0.08)
        cell.text = ""
        set_cell_margins(cell, 0, 0, 0, 0)
    repeat_header(accent.rows[0])
    keep_row_together(accent.rows[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    tag = p.add_run("ADVENTUREX 2026  ·  FROZEN CONCEPT v1.0")
    tag.bold = True
    tag.font.size = Pt(10)
    tag.font.color.rgb = RGBColor.from_string(ORANGE)
    set_east_asia(tag)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("RootScope")
    set_east_asia(run)

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(sub, "固定式根区灌溉舱")

    strap = doc.add_paragraph()
    strap.paragraph_format.space_after = Pt(5)
    add_inline(strap, "把看不见的根区给水，变成可见、可量、可追溯的端侧闭环。", size=10.5, color=DARK_BLUE, bold=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    add_inline(meta, "四人团队  |  72 小时执行方案  |  固定构型冻结版  |  2026-07-14", size=9.5, color=MUTED)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(2)
    shape = pic_p.add_run().add_picture(str(IMAGE), width=Inches(5.7))
    set_alt_text(shape, "RootScope 固定式根区灌溉舱三泵概念模型", "无轮桌面固定装置，包含透明薄层沙体卡匣、三台独立蠕动泵、称重储液瓶、RDK X5、STM32F407、固定相机、急停和干湿分区")

    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.space_after = Pt(8)
    add_inline(cap, "固定式三泵构型概念图 · 最终孔位与电气规格以实物测量和器件数据表为准")

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(callout, [BODY_WIDTH], BODY_WIDTH)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, 150, 220, 150, 220)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(cp, "最终拍板：1×RDK X5 上位智能 + 1×STM32F407 下位安全；三路固定深度、三泵物理隔离、称重停泵、固定视觉复核与 fail-closed 回执。", size=10.5, color=NAVY, bold=True)
    repeat_header(callout.rows[0])
    keep_row_together(callout.rows[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(7)
    add_inline(p, "设计目标：让 Expo 观众在 45 秒内亲手触发一次真实、可解释、可失败验证的物理闭环", size=9, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def table_widths(column_count: int, landscape: bool = False) -> tuple[list[int], int]:
    if landscape:
        total = 13968
        if column_count == 6:
            return [900, 2360, 2360, 2360, 2360, 3628], total
        return [total // column_count] * column_count, total
    presets = {
        2: [2600, 6760],
        3: [1900, 3650, 3810],
        4: [1300, 2450, 3150, 2460],
        5: [1200, 2040, 2040, 2040, 2040],
        6: [900, 1692, 1692, 1692, 1692, 1692],
    }
    return presets.get(column_count, [BODY_WIDTH // column_count] * column_count), BODY_WIDTH


def normalize_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = normalize_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(doc: Document, lines: list[str], landscape: bool) -> None:
    rows = [normalize_table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths, total = table_widths(col_count, landscape)
    if not landscape and rows[0][0] == "完成等级" and col_count == 3:
        widths = [2450, 3250, 3660]
    set_table_geometry(table, widths, total)
    set_repeat_table_header_and_rows(table)

    font_size = 7.6 if landscape and col_count >= 6 else (8.2 if col_count >= 4 else 9.0)
    for r_idx, data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(col_count):
            cell = row.cells[c_idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, PALE_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = data[c_idx] if c_idx < len(data) else ""
            add_inline(p, text, size=font_size, color=WHITE if r_idx == 0 else TEXT, bold=(r_idx == 0))
        row.height = None
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [BODY_WIDTH], BODY_WIDTH)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F4F7")
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    repeat_header(table.rows[0])
    keep_row_together(table.rows[0])
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor.from_string(NAVY)
        set_east_asia(run, latin="Consolas")
        if idx < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [110, 9250], BODY_WIDTH)
    set_cell_shading(table.cell(0, 0), ORANGE)
    set_cell_shading(table.cell(0, 1), PALE_BLUE)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 135, 160, 135, 160)
    table.cell(0, 0).text = ""
    repeat_header(table.rows[0])
    keep_row_together(table.rows[0])
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, size=10.5, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_body_image(doc: Document, alt: str, rel_path: str) -> None:
    path = (ROOT / rel_path).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(path), width=Inches(6.1))
    set_alt_text(shape, alt, alt)


def add_landscape_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    style_section(section, landscape=True, compact=True)
    add_header_footer(section)


def add_portrait_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    style_section(section)
    add_header_footer(section)


def create_decimal_numbering(doc: Document) -> int:
    """Create a fresh one-level decimal list so each Markdown block restarts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def parse_body(doc: Document, lines: list[str]) -> None:
    # Cover content ends immediately before the first numbered section.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    i = start
    landscape = False
    active_num_id: int | None = None
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if not line:
            active_num_id = None
            i += 1
            continue

        if line == "<!-- pagebreak -->":
            doc.add_page_break()
            active_num_id = None
            i += 1
            continue
        if line == "<!-- landscape:start -->":
            if not landscape:
                add_landscape_section(doc)
                landscape = True
            active_num_id = None
            i += 1
            continue
        if line == "<!-- landscape:end -->":
            if landscape:
                add_portrait_section(doc)
                landscape = False
            active_num_id = None
            i += 1
            continue
        if line.startswith("<!--") and line.endswith("-->"):
            active_num_id = None
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if not number_match:
            active_num_id = None

        if line.startswith("```"):
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip("\n"))
                i += 1
            add_code_block(doc, code)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            block = [line, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, block, landscape)
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^\)]+)\)", line)
        if image_match:
            add_body_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:].strip())
            set_paragraph_keep(p, keep_with_next=True)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip())
            set_paragraph_keep(p, keep_with_next=True)
            i += 1
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:].strip())
            i += 1
            continue
        if re.match(r"^-\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^-\s+", "", line))
            set_paragraph_keep(p, keep_together=True)
            i += 1
            continue
        if number_match:
            if active_num_id is None:
                active_num_id = create_decimal_numbering(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_num_id)
            add_inline(p, number_match.group(2))
            set_paragraph_keep(p, keep_together=True)
            i += 1
            continue
        if line.startswith("图 ") and "｜" in line:
            p = doc.add_paragraph(style="Caption")
            add_inline(p, line)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line)
        is_faq_question = line.startswith("**") and line.endswith("**")
        set_paragraph_keep(p, keep_with_next=is_faq_question, keep_together=is_faq_question)
        i += 1


def add_document_settings(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    option = OxmlElement("w:compatSetting")
    option.set(qn("w:name"), "compatibilityMode")
    option.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    option.set(qn("w:val"), "15")
    compat.append(option)


def remove_blank_trailing_paragraph(doc: Document) -> None:
    if not doc.paragraphs:
        return
    p = doc.paragraphs[-1]
    if not p.text.strip():
        parent = p._element.getparent()
        parent.remove(p._element)


def build() -> Path:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if not IMAGE.exists():
        raise FileNotFoundError(IMAGE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    install_styles(doc)
    add_document_settings(doc)
    doc.core_properties.title = "RootScope：AdventureX 2026 固定式根区灌溉舱执行方案"
    doc.core_properties.subject = "RDK X5 + STM32F407 固定式三泵根区灌溉舱的 72 小时工程方案"
    doc.core_properties.author = "RootScope 四人团队"
    doc.core_properties.keywords = "AdventureX, RootScope, RDK X5, STM32F407, 固定式灌溉舱, 端侧闭环"
    doc.core_properties.comments = "Frozen concept v1.0 · 2026-07-14"

    add_cover(doc)
    lines = SOURCE.read_text(encoding="utf-8-sig").splitlines()
    parse_body(doc, lines)
    remove_blank_trailing_paragraph(doc)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
